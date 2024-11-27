"""
Przykład zastosowania modułu docx
pip install python-docx
"""
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Create a new document
doc = Document()

# Title
doc.add_heading('Moduł JSON - Biblioteka Standardowa Python', level=1)

# Description
doc.add_paragraph(
    "Moduł json w Pythonie jest częścią biblioteki standardowej i służy do "
    "pracy z danymi w formacie JSON (JavaScript Object Notation). "
    "JSON jest popularnym formatem wymiany danych, szeroko stosowanym w API i aplikacjach webowych."
)

# Adding sections for functionalities
doc.add_heading('Funkcje podstawowe', level=2)

# Function: json.dumps()
doc.add_heading('json.dumps()', level=3)
doc.add_paragraph(
    "Funkcja json.dumps() konwertuje obiekt Python (np. słownik, listę) na ciąg JSON."
)
code = doc.add_paragraph()
code.style = doc.styles['Normal']
run = code.add_run('''
import json

# Przykład użycia json.dumps()
data = {"name": "Alice", "age": 25, "city": "New York"}
json_string = json.dumps(data, indent=4)
print(json_string)
''')
run.font.name = 'Courier New'
run.font.size = Pt(10)

# Function: json.loads()
doc.add_heading('json.loads()', level=3)
doc.add_paragraph(
    "Funkcja json.loads() konwertuje ciąg JSON na obiekt Python (np. słownik lub listę)."
)
code = doc.add_paragraph()
code.style = doc.styles['Normal']
run = code.add_run('''
import json

# Przykład użycia json.loads()
json_string = '{"name": "Alice", "age": 25, "city": "New York"}'
data = json.loads(json_string)
print(data["name"])  # Alice
''')
run.font.name = 'Courier New'
run.font.size = Pt(10)

# Function: json.dump()
doc.add_heading('json.dump()', level=3)
doc.add_paragraph(
    "Funkcja json.dump() zapisuje obiekt Python w pliku w formacie JSON."
)
code = doc.add_paragraph()
code.style = doc.styles['Normal']
run = code.add_run('''
import json

# Przykład użycia json.dump()
data = {"name": "Alice", "age": 25, "city": "New York"}
with open("data.json", "w") as file:
    json.dump(data, file, indent=4)
''')
run.font.name = 'Courier New'
run.font.size = Pt(10)

# Function: json.load()
doc.add_heading('json.load()', level=3)
doc.add_paragraph(
    "Funkcja json.load() odczytuje dane JSON z pliku i konwertuje je na obiekt Python."
)
code = doc.add_paragraph()
code.style = doc.styles['Normal']
run = code.add_run('''
import json

# Przykład użycia json.load()
with open("data.json", "r") as file:
    data = json.load(file)
print(data["city"])  # New York
''')
run.font.name = 'Courier New'
run.font.size = Pt(10)

# Saving the document
file_path = "d:\\Opis_modulu_JSON.docx"
doc.save(file_path)
